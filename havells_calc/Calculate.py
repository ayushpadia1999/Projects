from tkinter import *
import pandas as pd
from tabulate import tabulate
from docx import Document
import numpy as np
import os
import datetime
from datetime import timedelta
import time


root = Tk()
root.geometry("1260x320")
root.title("HAVELLS PRICE CALCULATOR")
#root.iconbitmap(r"C:\Users\ayush\OneDrive\Desktop\Havells\Havells_logo.ico")
photo = PhotoImage(file = 'Havells_logo.png')
root.iconphoto(False, photo)


load = 100
kw = 0
pole= 0
hp=0
F="160L"
B=3
S=1
IP=55
V=415
C='F'
IE=2
price=0
discount=0
dis_price=0
cat=0
num=0     
quantity = 1
df_price=pd.DataFrame()
df_quote=pd.DataFrame()
headers = [] 
df_price_temp=pd.DataFrame()
df_quote_temp=pd.DataFrame()
discount_list = []
load_list = []

def delete_row(event):
     row = dele.get()
     dele.delete(0,END)
     global df_price, df_quote,discount_list,load_list
     if(row.isnumeric()):
          row= int(row)
          df_price = df_price.drop(row, axis= 0 )  
          df_quote = df_quote.drop(row, axis =0  )
          #new code here
          discount_list.pop(row)
          load_list.pop(row)
          #print(df_quote)
          #print(df_price)
          quote.delete('1.0', END)
          quote.insert(INSERT,tabulate(df_quote, headers, tablefmt="pretty"))
          quote.insert(INSERT,"\n CATEGORY LIST:")
          quote.insert(INSERT,((df_price['Cat. '].values.tolist())))



def myClick(event):
     loadinge.delete(0,END)
     loadinge.insert(0, loading_calculate())
     pricee.delete(0,END)
     pricee.insert(0, price_calculate())
     dis_pricee.delete(0,END)
     dis_pricee.insert(0,0)
     #INSERT THE MATERIAL CODE HERE
     cate.delete(0,END)
     cate.insert(0,cat)
     quote.delete('1.0', END)

def myClick2(event):
     dis_pricee.delete(0,END)
     dis_pricee.insert(0,discounted_calculate())
     pole=int(polee.get())
     IE=int(IEe.get())
     global kw,hp,headers,df_quote_temp,quantity,F, discount_list
     quantity=quante.get()
     total_price=int(quantity)*discounted_calculate()
     data = [[kw,hp,int(polee.get()),
             "B" + str(int(Be.get())),"IE"+str(int(IEe.get())),F,"S"+str(int(Se.get())),Ce.get(),quantity,discounted_calculate(),total_price]]
     df_quote_temp=pd.DataFrame(data,columns = ["KW","HP","Pole","Mtg.",
                                           "Effy","Frame Size",
                                           "Duty","Insulation","Qty",
                                           "Net Price(Rs)","Total Price"])
     #New code here
     # discount_list.append(discounte.get())
     # print(discount_list)
     headers = ["KW","HP","Pole","Mtg.",
               "Effy","Frame Size",
               "Duty","Insulation","Qty",    
               "Net Price(Rs)","Total Price"]
     quote.delete('1.0', END)
     quote.insert(INSERT,tabulate(df_quote_temp, headers, tablefmt="pretty" ))

def myClick3(event):
     kwe.delete(0,END)
     polee.delete(0,END)
     hpe.delete(0,END)
     Fe.delete(0,END)
     Fe.insert(0,"160L")
     Be.delete(0,END)
     Be.insert(0,"3")
     Se.delete(0,END)
     Se.insert(0,"1")
     IPe.delete(0,END)
     IPe.insert(0,"55")
     Ve.delete(0,END)
     Ve.insert(0,"415")
     Ce.delete(0,END)
     Ce.insert(0,"F")
     IEe.delete(0,END)
     IEe.insert(0,"2")
     pricee.delete(0,END)
     pricee.insert(0,0)
     loadinge.delete(0,END)
     loadinge.insert(0,0)
     dis_pricee.delete(0,END)
     dis_pricee.insert(0,0)
     discounte.delete(0,END)
     discounte.insert(0,0)
     cate.delete(0,END)
     cate.insert(0,0)
     load = 100
     global kw 
     kw = 0
     global pole
     pole= 0
     global hp
     hp = 0
     global F
     F="160L"
     global B
     B=3
     global S
     S=1
     global IP
     IP=55
     global V
     V=415
     global C
     C ='F'
     global IE
     IE  =2
     global price
     price=0
     global discount
     discount=0
     global dis_price
     dis_price=0
     global cat
     cat=0
     global num
     num =0     
     global df_price_temp,df_quote_temp,df_price,df_quote
     df_price_temp=pd.DataFrame()
     df_quote_temp=pd.DataFrame()
     df_quote=pd.DataFrame()
     df_price=pd.DataFrame()
     # print(df_price)
     quote.delete('1.0', END)
     #quote.insert(INSERT,tabulate(df_quote, headers, tablefmt="pretty" , showindex=False))
     #quote.insert(INSERT,tabulate(df_quote, headers, tablefmt="pretty" , showindex=False))

#GENERATE QUOTE BUTTON
def myClick4(event):
     #out=Document('Test.docx')
     #doc_table = out.tables[0]
     files = ['Test.docx', 'OFFER_FORMAT.docx']

     merged_document = Document()

     for index, file in enumerate(files):
          sub_doc = Document(file)

        # Don't add a page break if you've reached the last file.
          if index < len(files)-1:
               sub_doc.add_page_break()

          for element in sub_doc.element.body:
               merged_document.element.body.append(element)

     merged_document.save('OFFER_FORMAT.docx') 

     df_price['Cat. '].to_clipboard( sep = '\n' ,index= False , header =False)
     doc = Document()
     #ADD FORMAT
     offer = Document('OFFER_FORMAT.docx')
     ref_num = str(offer.paragraphs[0].runs[0].text)
     # print(ref_num[27:34])
     ref_num = int(ref_num[27:34])

     x = datetime.datetime.now()
     ref_num = ref_num + 1
     offer.save('OFFER_FORMAT.docx')
     ref_num = (ref_num)%100000
     DATE = doc.add_paragraph('REF: AA/Q/' + str(x.year) + "-" + str(x.year+1) + "/HMA-SL-S" + str(ref_num) + "                                                                "  + 
     "DATE: " + str(x.day )+ "/" + str(x.month) + "/" + str(x.year) )
     DATE.add_run("\n \n Dear Sirs, \n\n With reference to your enquiry, we are pleased to offer our best rates for Havells Make Motors hereunder: ")
     #FINISH FORMAT
     t = doc.add_table(df_quote.shape[0]+1, df_quote.shape[1])
     t.style = 'Table Grid'
     # add the header rows.
     for j in range(df_quote.shape[-1]):
          t.cell(0,j).text = df_quote.columns[j]
     # add the rest of the data frame
     for i in range(df_quote.shape[0]):
          for j in range(df_quote.shape[-1]):
             t.cell(i+1,j).text = str(df_quote.values[i,j])
     row = t.rows[0]
     for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
     #Add Format
     #NEW CODE TO BE ADDED
     f = open("Record.txt","a+")
     f.write('\nREF: AA/Q/' + str(x.year) + "-" + str(x.year+1) + "/HMA-SL-" + str(ref_num) + "                                                                                    "  + 
     "DATE: " + str(x.day )+ "/" + str(x.month) + "/" + str(x.year)+'\n' )
     f.write(tabulate(df_quote.assign(Dicount = discount_list , Loading = load_list),headers + ['Discount %'] + ['Loading %'], tablefmt="pretty" , showindex=False))
     f.close()
     # print(discount_list)
     Last = doc.add_paragraph()
     run = Last.add_run("\n\n\nOther Terms & Conditions:")
     run.underline = True
     run.bold = True
     Last.add_run("\n1. Prices                                :  F.O.R. Your Plant") 
     Last.add_run("\n2. Taxes                                 :  GST will be charged extra @18%")
     Last.add_run("\n3. Payment                           : 100% Within 30 Days")
     x = (x+timedelta(days=15))
     Last.add_run("\n4. Validity                             :  " + str(x.day )+ "/" + str(x.month) + "/" + str(x.year) )
     Last.add_run("\n5. Delivery                            :") 

     Last.add_run("\n \n \n We keep ourselves open for any further technical or commercial clarification needed in this regard. ")
     Last.add_run("\nThanking you and assuring you best of our services at all times. \n Yours Faithfully, \n")
     Last.add_run("Amit Padia \n M-9433227876\n").bold =True
     run = Last.add_run('AMIT & ASSOCIATES \n')
     font = run.font
     from docx.shared import Pt
     font.name = 'Engravers MT'
     font.size = Pt(16)
     Last.add_run('9, Ezra Street, 2nd Floor, Kolkata-700001. \nPh: 033-22253967/9830874211 \nEmail: ').bold=True
     from docx.shared import RGBColor
     run = Last.add_run('amitpadia73@gmail.com ')
     run.bold = True
     run.underline = True
     font = run.font
     font.color.rgb = RGBColor(51,102,187)
     Last.add_run("and").bold = True 
     run = Last.add_run(' amitpadia@rediffmail.com ')
     run.bold = True
     run.underline = True
     font = run.font
     font.color.rgb = RGBColor(51,102,187)
     Last.add_run("\nGST No. 19AFHPP9819G1ZV").bold = True
     #FInish Format
     doc.save('Test.docx')
     os.startfile('Test.docx')
     #f = open('Test.docx')    
     #time.sleep(5)
     #tmp=0
    #while(tmp==0):
          #try:
               #docx = Document('Test.docx')
               #docx.save('Test.docx')
               #print("Yay")
               #tmp=1
          #except:
               #tmp=0
               #print("Wait")

     #print("CLOSE")
     #time.sleep(5)
      


#CONFIRM BUTTON
def myClick5(event):
     global df_price,df_quote,df_price_temp,discount_list,load_list
     df_price = df_price.append(df_price_temp , ignore_index=True)
     df_quote = df_quote.append(df_quote_temp , ignore_index=True)
     quote.delete('1.0', END)
     quote.insert(INSERT,tabulate(df_quote, headers, tablefmt="pretty" ))
     quote.insert(INSERT,"\n CATEGORY LIST:")
     #quote.insert(INSERT,((df_price['Cat. ']).to_string(index=False)))
     quote.insert(INSERT,((df_price['Cat. '].values.tolist())))
     discount_list.append(discounte.get())
     load_list.append(loadinge.get())
     #print(df_quote)
 
#CANCEL BUTTON
def myClick6(event):
     kwe.delete(0,END)
     polee.delete(0,END)
     hpe.delete(0,END)
     Fe.delete(0,END)
     Fe.insert(0,"160L")
     Be.delete(0,END)
     Be.insert(0,"3")
     Se.delete(0,END)
     Se.insert(0,"1")
     IPe.delete(0,END)
     IPe.insert(0,"55")
     Ve.delete(0,END)
     Ve.insert(0,"415")
     Ce.delete(0,END)
     Ce.insert(0,"F")
     IEe.delete(0,END)
     IEe.insert(0,"2")
     pricee.delete(0,END)
     pricee.insert(0,0)
     loadinge.delete(0,END)
     loadinge.insert(0,0)
     dis_pricee.delete(0,END)
     dis_pricee.insert(0,0)
     discounte.delete(0,END)
     discounte.insert(0,0)
     cate.delete(0,END)
     cate.insert(0,0)
     load = 100
     global kw 
     kw = 0
     global pole
     pole= 0
     global hp
     hp = 0
     global F
     F="160L"
     global B
     B=3
     global S
     S=1
     global IP
     IP=55
     global V
     V=415
     global C
     C ='F'
     global IE
     IE  =2
     global price
     price=0
     global discount
     discount=0
     global dis_price
     dis_price=0
     global cat
     cat=0
     global num
     num =0     
     global df_price_temp,df_quote_temp
     df_price_temp=pd.DataFrame()
     df_quote_temp=pd.DataFrame()
     quote.delete('1.0', END)
     quote.insert(INSERT,tabulate(df_quote, headers, tablefmt="pretty" ))
     quote.insert(INSERT,"\n CATEGORY LIST:")
     quote.insert(INSERT,((df_price['Cat. '].values.tolist())))
     #print(df_quote)



kwe = Entry(root)
polee = Entry(root)
hpe = Entry(root)
Fe = Entry(root)
Be = Entry(root)
Se = Entry(root)
IPe = Entry(root)
Ve = Entry(root)
Ce = Entry(root)
IEe= Entry(root)

kwe.grid(row=3,column=0)
polee.grid(row=3,column=1)
hpe.grid(row=3,column=2)
Fe.grid(row=3,column=3)
Be.grid(row=3,column=4)
Se.grid(row=3,column=5)
IPe.grid(row=3,column=6)
Ve.grid(row=3,column=7)
Ce.grid(row=3,column=8)
IEe.grid(row=3,column=9)

Fe.insert(0,"160L")
Be.insert(0,"3")
Se.insert(0,"1")
IPe.insert(0,"55")
Ve.insert(0,"415")
Ce.insert(0,"F")
IEe.insert(0,"2")

myButton = Button(root, text="Calculate Price")
myButton.bind("<Return>",myClick)
myButton.bind("<Button-1>",myClick)
myButton.grid(row=4,column=3,columnspan=2, padx = 50 , pady =20)            
#percentage increase = loading
loadl = Label(root, text="LOADING CALCULATED:")
loadl.grid(row= 5,column =0)
loadinge = Entry(root)
loadinge.grid(row= 5 ,column = 1)
loadinge.insert(0,0)
#price calculated
pricel = Label(root, text="PRICE CALCULATED:")
pricel.grid(row= 6,column =0)
pricee = Entry(root)
pricee.grid(row= 6 ,column = 1)
pricee.insert(0,0)
 #discount fetch
discountl = Label(root, text="ENTER DISCOUNT:")
discountl.grid(row= 7,column =0)
discounte = Entry(root)
discounte.grid(row= 7 ,column = 1)
discounte.insert(0,0)
#discounted price
dis_pricel = Label(root, text="DISCOUNTED PRICE:")
dis_pricel.grid(row= 8, column=0)
dis_pricee = Entry(root)
dis_pricee.grid(row=8 , column=1 )
dis_pricee.insert(0,0)
#Quantity 
quantl = Label(root, text="QUANTITY")
quantl.grid(row= 9, column=0)
quante = Entry(root)
quante.grid(row=9 , column=1 )
quante.insert(0,1)
#category
catl = Label(root, text="MATERIAL CODE:")
catl.grid(row= 10, column=0)
cate = Entry(root)
cate.grid(row=10 , column=1 )
cate.insert(0,0)

def loading_calculate():
     global load
     load = 100
     if(int(Be.get())!=3):
          load=load+5
          #print(load)
     if(int(Se.get())!=1):
          load=load+7
          #print(load)
     if(int(Ve.get())!=415):
          load=load+3
          #print(load)
     if(Ce.get()!='F'):
          load=load+10
          #print(load)
     if(int(IPe.get())==56):
          load=load+5
          #print(load)
     if(int(IPe.get())==65 or int(IPe.get())==66):
          load=load+10
          #print(load)
     load_local=str(load-100)
     return load_local+"%"
    
def price_calculate():
     global load
     df=pd.read_csv('HavellsPriceSheet.csv')
     df.drop(df.columns[0], axis = 1, inplace = True)
     global kw,pole,hp,IE,cat,price,df_price_temp,F
     pole=int(polee.get())
     IE=int(IEe.get())
     if((hpe.get()!='') and (kwe.get()!='')):
          quote.delete('1.0', END)
          quote.insert(INSERT,"                                           DONT ENTER ENTER BOTH KW AND HP ")
     elif(hpe.get()!=''):
          hp=float(hpe.get())
          df_price_temp = df.loc[((df["Pole"]==pole)&(df["HP"]==hp))&(df["IE No."]==IE)]
     elif(kwe.get()!=''):
          kw=float(kwe.get())
          df_price_temp = df.loc[((df["Pole"]==pole)&(df["kW"]==kw))&(df["IE No."]==IE)]
     else:
          quote.delete('1.0', END)
          quote.insert(INSERT,"                                           ENTER KW OR HP ")
     #print(type(pole))    
     df_price_temp = df_price_temp.reset_index(drop=True)
     #print(df_price_temp)
     price = df_price_temp.iloc[0]["List Price"]
     cat = df_price_temp.iloc[0]["Cat. "]
     kw = df_price_temp.iloc[0]["kW"]
     hp = df_price_temp.iloc[0]["HP"]
     F = df_price_temp.iloc[0]["Frame"]
     #load = float(loadinge.get())
     load=load/100
     #print(load)
     return price*load

def discounted_calculate():
    discount = float(discounte.get())
    price = float(pricee.get())
    dis_price = price*(100-discount)/100
    return round(dis_price)

myButton2= Button(root, text="Calculate Discounted Price")
myButton2.bind("<Return>",myClick2)
myButton2.bind("<Button-1>",myClick2)

myButton5= Button(root, text="Confirm")
myButton5.bind("<Return>",myClick5)
myButton5.bind("<Button-1>",myClick5)

myButton6= Button(root, text="Cancel")
myButton6.bind("<Return>",myClick6)
myButton6.bind("<Button-1>",myClick6)

myButton4 = Button(root, text="Generate Quote")
myButton4.bind("<Return>",myClick4)
myButton4.bind("<Button-1>",myClick4)

myButton3 = Button(root, text="Reset Entries")
myButton3.bind("<Return>",myClick3)
myButton3.bind("<Button-1>",myClick3)

myButton2.grid(row=4,column=4,columnspan=2 )
myButton5.grid(row=4,column=5,columnspan=2)
myButton6.grid(row=4,column=6)
myButton4.grid(row=4,column=7)
myButton3.grid(row=4,column=8)   

dell = Label(root, text="DELETE ENTRY:")
dell.grid(row= 4, column=0)
dele = Entry(root)
dele.grid(row=4 , column=1 )
dele.bind("<Return>", delete_row)

myLabel1 = Label(root,text="HAVELLS PRICE DETAILS")
myLabel2 = Label(root,text="Please Enter the Product Details: ")
kwLabel = Label(root,text="Kilowatt ")
polelabel = Label(root,text="Pole ")
hplabel = Label(root,text="Horse Power ")
FLabel = Label(root,text="Frame ")
BLabel = Label(root,text="Mounting ")
SLabel = Label(root,text="Duty ")
IPLabel = Label(root,text="IP No. ")
VLabel = Label(root,text="Voltage ")
CLabel = Label(root,text="Class")
IELabel = Label(root,text="IE No.")
quote = Text(root,height = 9, width = 120)

myLabel1.grid(row=0,column=0,columnspan=10)
myLabel2.grid(row=1,column=0,columnspan=10)
kwLabel.grid(row=2,column=0)
polelabel.grid(row=2,column=1)
hplabel.grid(row=2,column=2)
FLabel.grid(row=2,column=3)
BLabel.grid(row=2,column=4)
SLabel.grid(row=2,column=5)
IPLabel.grid(row=2,column=6)
VLabel.grid(row=2,column=7)
CLabel.grid(row=2,column=8)
IELabel.grid(row=2,column=9)
quote.grid(row=5, column=2,rowspan=5,columnspan=11)
quote.insert(INSERT,"                                                     QUOTATION FOR MOTOR ")
root.mainloop()

